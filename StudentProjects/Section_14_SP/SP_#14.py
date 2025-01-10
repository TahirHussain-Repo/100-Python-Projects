import os
import docx

folder_path = 'word_documents'
output_file = 'word_documents/merged_documents.docx'

filenames = os.listdir(folder_path)

filepaths = [os.path.join(folder_path, filename) for filename in filenames if filename.endswith('docx')]

merged_document = docx.Document()

for filepath in filepaths:
    doc = docx.Document(filepath)

for para in doc.paragraphs:
    new_para = merged_document.add_paragraph()
    new_para._element.addprevious(para._element)

merged_document.save(output_file)



