import docx

docpath1 = 'panda1.docx'
docpath2 = 'panda2.docx'

doc1 = docx.Document(docpath1)
doc2 = docx.Document(docpath2)

para = doc2.paragraphs[0]

#new_paragraph = doc1.add_paragraph(text)
paragraphs = doc1.paragraphs
paragraphs[1]._element.addnext(para._element)
doc1.save('panda.docx')
